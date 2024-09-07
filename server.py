from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
from flask_cors import CORS
import google.generativeai as genai
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app) #Cross-Origin Resource Sharing


UPLOAD_FOLDER = 'uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)



genai.configure(api_key=os.getenv('YOUR_GEMINI_API_KEY'))
model = genai.GenerativeModel("gemini-1.5-flash")

doc_name = 'new_convo.docx'
doc = Document()

if os.path.exists(doc_name):
    doc = Document(doc_name)


doc2 = Document('docs/Employee Policy.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/HR Policy 1.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/Remote Work Policy.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/IT POLICY MANUAL.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/HR POLICY MANUAL.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/Company Healthcare Policy.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/Company Transportation Policy.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])

doc2 = Document('docs/Company Workplace Security Policy.docx')
doc2="".join([paragraph.text for paragraph in doc2.paragraphs])


doc2_response = model.generate_content("it should have all info from this para but summarize it\n\n\n"+doc2).text 
print("completed_summaries")

# Convert PDF to text
def pdf_to_text(file_path):
    pdf_reader = PdfReader(file_path)
    text = "".join(page.extract_text() for page in pdf_reader.pages)
    return text

# Convert DOCX to text
def docx_to_text(file_path):
    document = Document(file_path)
    return "\n".join([paragraph.text for paragraph in document.paragraphs])

# Handle text response with entire conversation history
def handle_text_input(input_text):
    
    doc.add_paragraph(f"User: {input_text}") #append new input

    #Combine entire conversation 
    conversation_history = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    full_text = f"{conversation_history}\n\n'only response for this, the above is the history,dont give any formating  '\n\n\nUser: {input_text}"
    
    #chatgpt_api
    response = model.generate_content(full_text).text 
    
    #Save the response 
    doc.add_paragraph(f"Bot: {response}")
    doc.add_paragraph(doc2_response)
    
    doc.save(doc_name)
    return response

@app.route('/', methods=['GET'])
def test():
    doc = Document(doc_name)
    conversation_history = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    #return jsonify({"message":doc2_response})  #use this for testing
    return jsonify({"message":conversation_history}) #cant use this for testing because there is no file created


@app.route('/bot', methods=['POST'])
def process_request():
    text_input = request.form.get('text', '')
    file = request.files.get('file')
    r_format = request.form.get('format', '').lower()

    if not text_input and not file:
        return jsonify({"error": "No input provided"}), 400

    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        if r_format == "pdf":
            extracted_text = pdf_to_text(file_path)
        elif r_format == "docx":
            extracted_text = docx_to_text(file_path)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
        
        response = handle_text_input(extracted_text)
        return jsonify({"status": "File processed successfully", "text": response}), 200

    if text_input:
        response = handle_text_input(text_input)
        return jsonify({"status": "Query processed successfully", "text": response}), 200
    


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)